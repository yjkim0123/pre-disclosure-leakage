#!/usr/bin/env python3
"""Create proper Title Page as Word doc"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Non-Directional Pre-Disclosure Drift Under Fair Disclosure:\nEvidence from Korean Earnings Announcements')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
p.paragraph_format.space_after = Pt(24)

# Author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Yongjun Kim')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
p.paragraph_format.space_after = Pt(6)

# Affiliation
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Department of Software, Ajou University\nSuwon, South Korea\nyjkim0123@ajou.ac.kr')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(24)

# Corresponding author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Corresponding Author: Yongjun Kim\nEmail: yjkim0123@ajou.ac.kr\nORCID: 0000-0003-4234-4883')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(36)

# Acknowledgements
p = doc.add_paragraph()
run = p.add_run('Acknowledgements')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('The author thanks the DART electronic disclosure system for providing publicly accessible corporate filing data.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(12)

# Funding
p = doc.add_paragraph()
run = p.add_run('Funding')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
p.paragraph_format.space_after = Pt(12)

# Declaration
p = doc.add_paragraph()
run = p.add_run('Declaration of Competing Interest')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph()
run = p.add_run('The author declares no competing financial interests or personal relationships that could have influenced the work reported in this paper.')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.save('/Users/yongjun_kim/.openclaw/workspace/dart_title_page.docx')
print("✅ Title page created!")
