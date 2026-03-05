#!/usr/bin/env python3
"""Fix pandoc-generated docx formatting"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import copy

doc = Document('/Users/yongjun_kim/.openclaw/workspace/dart_manuscript_blind_v2.docx')

# Fix font and spacing for all paragraphs
for para in doc.paragraphs:
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    # Fix spacing
    fmt = para.paragraph_format
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.5

# Fix heading styles  
for para in doc.paragraphs:
    if para.style.name.startswith('Heading'):
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            if '1' in para.style.name:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)

# Fix tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                para.paragraph_format.space_after = Pt(2)
                para.paragraph_format.line_spacing = 1.0

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

doc.save('/Users/yongjun_kim/.openclaw/workspace/dart_manuscript_blind_v2.docx')
print("✅ Done! Formatting fixed.")
