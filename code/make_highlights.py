#!/usr/bin/env python3
"""Create Highlights as Word doc"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

p = doc.add_paragraph()
run = p.add_run('Highlights')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.bold = True
p.paragraph_format.space_after = Pt(12)

highlights = [
    'Pre-disclosure abnormal returns exist but do not predict the direction of actual earnings surprises (49.0% accuracy, no better than chance).',
    'Non-parametric tests (sign test, Wilcoxon) fail to confirm systematic directional drift in pre-disclosure returns.',
    'Post-disclosure Day +1 returns respond sharply and correctly to earnings surprises (+1.13 p.p. spread, p<0.001, 60.6% direction accuracy).',
    '96.5% of DART earnings filings occur after market close, making Day 0 effectively pre-disclosure.',
    'Results are robust to firm-clustered and month-clustered standard errors.',
]

for h in highlights:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(h)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5

doc.save('/Users/yongjun_kim/.openclaw/workspace/dart_highlights.docx')
print("✅ Highlights created!")
